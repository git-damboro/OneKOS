from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "OneKOS-复赛OnePage.md"
OUTPUT = ROOT / "docs" / "OneKOS-复赛OnePage.docx"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margin(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_rich_text(paragraph, text):
    pieces = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(72, 86, 118)
        else:
            paragraph.add_run(piece)


def apply_body_font(paragraph, size=9.5, color="253047"):
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.12


def add_table(document, rows):
    headers = [item.strip() for item in rows[0].strip().strip("|").split("|")]
    data_rows = rows[2:]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header_cells = table.rows[0].cells
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_tr_pr.append(repeat_header)
    for index, value in enumerate(headers):
        header_cells[index].text = value
        set_cell_fill(header_cells[index], "273A70")
        set_cell_margin(header_cells[index])
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in header_cells[index].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
    for row_index, line in enumerate(data_rows):
        values = [item.strip() for item in line.strip().strip("|").split("|")]
        cells = table.add_row().cells
        row_properties = table.rows[-1]._tr.get_or_add_trPr()
        no_split = OxmlElement("w:cantSplit")
        no_split.set(qn("w:val"), "true")
        row_properties.append(no_split)
        for index, value in enumerate(values):
            cell = cells[index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_rich_text(paragraph, value)
            apply_body_font(paragraph, 8.5)
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 1:
                set_cell_fill(cell, "F5F7FC")
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def build_document():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    for style_name, size, color in (("Title", 22, "25386D"), ("Heading 1", 15, "25386D"), ("Heading 2", 11.5, "586BF6")):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "千面·OneKOS｜AI 驱动的个性化 KOS 内容增长操作系统"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    apply_body_font(header, 7.5, "7A849A")

    footer = section.footer.paragraphs[0]
    footer.text = "AI 先锋未来人才大赛｜复赛方案初稿｜模拟数据演示"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_body_font(footer, 7.5, "7A849A")

    index = 0
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.right_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(45, 58, 88)
                p.style = doc.styles["Normal"]
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[index + 1]):
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, table_lines)
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)
            p.add_run(line[2:])
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.add_run(line[3:])
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            p.add_run(line[4:])
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.45)
            p.paragraph_format.right_indent = Cm(0.45)
            p.paragraph_format.space_after = Pt(7)
            add_rich_text(p, line[2:])
            apply_body_font(p, 11, "354A83")
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\. ", "", line))
            apply_body_font(p)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:])
            apply_body_font(p)
        elif line.strip():
            p = doc.add_paragraph()
            add_rich_text(p, line)
            apply_body_font(p)
        index += 1

    doc.core_properties.title = "千面·OneKOS 复赛方案初稿 OnePage"
    doc.core_properties.subject = "AI 驱动的个性化 KOS 内容增长操作系统"
    doc.core_properties.author = "千面·OneKOS"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
